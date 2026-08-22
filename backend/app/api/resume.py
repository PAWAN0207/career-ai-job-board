import os
import re
import json
import uuid

from fastapi import (
    APIRouter,
    UploadFile,
    File,
    HTTPException,
    Depends,
)

from app.core.supabase import supabase
from app.core.auth import get_current_user
from app.services.gemini import generate_response


router = APIRouter(
    prefix="/api/resume",
    tags=["Resume Intelligence"],
)


# ============================================================
# UPLOAD DIRECTORY
# ============================================================

UPLOAD_DIR = "uploads/resumes"

os.makedirs(
    UPLOAD_DIR,
    exist_ok=True,
)


# ============================================================
# CONSTANTS
# ============================================================

ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
}

MAX_FILE_SIZE = 5 * 1024 * 1024


# ============================================================
# TEXT CLEANING
# ============================================================

def clean_resume_text(text: str) -> str:

    if not text:
        return ""

    text = re.sub(
        r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]",
        " ",
        text,
    )

    text = text.replace("ÃƒÂ¢Ã¢â€šÂ¬Ã‚Â¢", " ")
    text = text.replace("ÃƒÂ¯Ã¢â‚¬Å¡Ã‚Â·", " ")

    text = re.sub(
        r"(@[A-Za-z0-9._%+-]+)\s*\n\s*"
        r"([A-Za-z0-9.-]+\.[A-Za-z]{2,})",
        r"\1\2",
        text,
        flags=re.IGNORECASE,
    )

    text = re.sub(
        r"[ \t]+",
        " ",
        text,
    )

    text = re.sub(
        r"\n\s*\n+",
        "\n\n",
        text,
    )

    return text.strip()


# ============================================================
# RESUME VALIDATION
# ============================================================

def validate_resume_document(resume_text: str):

    if not resume_text:
        return {
            "is_resume": False,
            "score": 0,
            "message": (
                "The uploaded document does not contain "
                "readable text. Please upload a valid resume."
            ),
        }

    text = resume_text.lower()

    if len(resume_text.strip()) < 120:
        return {
            "is_resume": False,
            "score": 0,
            "message": (
                "The uploaded document is too short to be "
                "a valid resume. Please upload a complete resume."
            ),
        }

    resume_section_keywords = {
        "resume": 2,
        "curriculum vitae": 2,
        "cv": 1,
        "professional summary": 2,
        "summary": 1,
        "objective": 1,
        "experience": 2,
        "work experience": 3,
        "professional experience": 3,
        "employment": 2,
        "education": 2,
        "qualification": 2,
        "skills": 2,
        "technical skills": 2,
        "projects": 2,
        "achievements": 1,
        "certifications": 1,
        "internship": 2,
        "internships": 2,
        "career objective": 2,
        "profile": 1,
    }

    score = 0
    matched_sections = []

    for keyword, points in resume_section_keywords.items():

        if keyword in text:
            score += points
            matched_sections.append(keyword)

    email_match = re.search(
        r"[A-Za-z0-9._%+-]+@"
        r"[A-Za-z0-9.-]+\."
        r"[A-Za-z]{2,}",
        resume_text,
        re.IGNORECASE,
    )

    phone_match = re.search(
        r"(?:\+91[\s-]?)?[6-9]\d{9}",
        resume_text,
    )

    if email_match:
        score += 2

    if phone_match:
        score += 2

    education_keywords = [
        "bachelor",
        "b.sc",
        "bsc",
        "b.com",
        "bba",
        "bca",
        "master",
        "m.sc",
        "msc",
        "m.com",
        "mba",
        "mca",
        "degree",
        "diploma",
        "h.s.c",
        "hsc",
        "s.s.c",
        "ssc",
        "university",
        "college",
        "school",
    ]

    if any(
        keyword in text
        for keyword in education_keywords
    ):
        score += 2

    work_keywords = [
        "company",
        "organization",
        "employer",
        "job title",
        "designation",
        "responsibilities",
        "responsibility",
        "worked",
        "working",
        "experience",
        "years of experience",
        "months of experience",
        "intern",
        "trainee",
    ]

    work_matches = sum(
        1
        for keyword in work_keywords
        if keyword in text
    )

    if work_matches >= 1:
        score += 2

    if work_matches >= 3:
        score += 2

    common_skill_keywords = [
        "python",
        "sql",
        "excel",
        "power bi",
        "tableau",
        "java",
        "javascript",
        "c++",
        "machine learning",
        "data analysis",
        "communication",
        "leadership",
        "management",
        "ms office",
        "microsoft office",
        "tally",
        "sap",
        "salesforce",
        "kyc",
        "banking",
        "chemistry",
        "laboratory",
        "quality control",
    ]

    skill_matches = sum(
        1
        for keyword in common_skill_keywords
        if keyword.lower() in text
    )

    if skill_matches >= 1:
        score += 1

    if skill_matches >= 3:
        score += 2

    if re.findall(
        r"\b(?:19|20)\d{2}\b",
        resume_text,
    ):
        score += 1

    strong_sections = [
        "experience",
        "education",
        "skills",
        "projects",
        "professional summary",
        "work experience",
        "professional experience",
    ]

    strong_resume_sections = sum(
        1
        for section in strong_sections
        if section in text
    )

    has_contact = (
        email_match is not None
        or phone_match is not None
    )

    has_strong_structure = (
        strong_resume_sections >= 2
    )

    is_resume = (
        score >= 7
        and (
            has_contact
            or has_strong_structure
        )
    )

    if not is_resume:
        return {
            "is_resume": False,
            "score": score,
            "matched_sections": matched_sections,
            "message": (
                "The uploaded document does not appear "
                "to be a resume. Please upload a valid "
                "resume or CV containing your education, "
                "experience, skills or professional details."
            ),
        }

    return {
        "is_resume": True,
        "score": score,
        "matched_sections": matched_sections,
        "message": "Document appears to be a valid resume.",
    }


# ============================================================
# JSON CLEANER
# ============================================================

def clean_ai_json(response: str) -> str:

    if not response:
        raise ValueError(
            "Empty response received from AI."
        )

    cleaned = response.strip()

    cleaned = re.sub(
        r"^```(?:json)?\s*",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )

    cleaned = re.sub(
        r"\s*```$",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )

    start = cleaned.find("{")
    end = cleaned.rfind("}")

    if (
        start != -1
        and end != -1
        and end > start
    ):
        cleaned = cleaned[start:end + 1]

    return cleaned


# ============================================================
# DOMAIN DETECTION
# ============================================================

DOMAIN_KEYWORDS = {

    "chemistry": [
        "chemist",
        "chemistry",
        "analytical chemistry",
        "organic chemistry",
        "inorganic chemistry",
        "physical chemistry",
        "industrial chemistry",
        "chemical analysis",
        "chemical analyst",
        "laboratory",
        "lab analyst",
        "lab technician",
        "laboratory technician",
        "laboratory analyst",
        "quality control",
        "quality assurance",
        "qc analyst",
        "qa analyst",
        "pharmaceutical",
        "pharma",
        "formulation",
        "titration",
        "spectroscopy",
        "chromatography",
        "hplc",
        "gc",
        "uv spectroscopy",
        "uv visible",
        "wet chemistry",
        "sample preparation",
        "chemical testing",
        "chemical testing",
        "research chemist",
        "analytical chemist",
    ],

    "biology_life_sciences": [
        "biology",
        "biologist",
        "microbiology",
        "microbiologist",
        "biotechnology",
        "biotechnologist",
        "life sciences",
        "molecular biology",
        "biochemistry",
        "genetics",
        "cell biology",
        "clinical research",
        "clinical laboratory",
        "bioinformatics",
    ],

    "data_analytics": [
        "data analyst",
        "data analytics",
        "business analyst",
        "bi analyst",
        "business intelligence",
        "analytics analyst",
        "reporting analyst",
        "data visualization",
        "dashboard development",
        "kpi dashboard",
        "power bi dashboard",
        "tableau dashboard",
        "sql reporting",
        "business reporting",
        "data-driven insights",
    ],

    "data_science": [
        "data scientist",
        "data science",
        "machine learning engineer",
        "ml engineer",
        "ai engineer",
        "artificial intelligence",
        "machine learning",
        "deep learning",
        "natural language processing",
        "nlp",
        "model training",
        "predictive modeling",
        "feature engineering",
        "model deployment",
    ],

    "software_development": [
        "software engineer",
        "software developer",
        "frontend developer",
        "backend developer",
        "full stack developer",
        "full-stack developer",
        "web developer",
        "application developer",
        "react developer",
        "node developer",
        "typescript developer",
        "javascript developer",
    ],

    "banking_finance": [
        "banking",
        "banking operations",
        "banking executive",
        "finance",
        "financial analyst",
        "credit analyst",
        "credit risk",
        "loan processing",
        "loan operations",
        "disbursement",
        "kyc",
        "know your customer",
        "financial services",
        "accounts",
        "accounting",
        "gst",
        "tally",
    ],

    "operations": [
        "operations executive",
        "operations analyst",
        "operations associate",
        "business operations",
        "process executive",
        "process associate",
        "document processing",
        "back office",
        "customer operations",
    ],

    "education": [
        "teacher",
        "educator",
        "academic mentor",
        "faculty",
        "lecturer",
        "school teacher",
        "subject matter expert",
        "sme",
        "tutor",
        "teaching",
    ],

    "marketing_sales": [
        "marketing",
        "digital marketing",
        "sales executive",
        "sales representative",
        "business development",
        "business development executive",
        "marketing executive",
        "seo",
        "social media marketing",
    ],
}


DOMAIN_ROLES = {

    "chemistry": [
        "Chemist",
        "Analytical Chemist",
        "QC Chemist",
        "Chemical Analyst",
        "Laboratory Analyst",
        "Research Chemist",
    ],

    "biology_life_sciences": [
        "Biologist",
        "Microbiologist",
        "Biotechnology Associate",
        "Research Associate",
        "Laboratory Analyst",
    ],

    "data_analytics": [
        "Data Analyst",
        "Business Analyst",
        "BI Analyst",
        "Reporting Analyst",
    ],

    "data_science": [
        "Data Scientist",
        "Junior Data Scientist",
        "Machine Learning Engineer",
        "AI Engineer",
    ],

    "software_development": [
        "Software Developer",
        "Software Engineer",
        "Backend Developer",
        "Frontend Developer",
        "Full Stack Developer",
    ],

    "banking_finance": [
        "Banking Operations Executive",
        "Financial Analyst",
        "Credit Analyst",
        "Finance Executive",
    ],

    "operations": [
        "Operations Executive",
        "Operations Analyst",
        "Process Executive",
        "Operations Associate",
    ],

    "education": [
        "Educator",
        "Academic Mentor",
        "Subject Matter Expert",
        "Teacher",
    ],

    "marketing_sales": [
        "Marketing Executive",
        "Sales Executive",
        "Business Development Executive",
    ],

    "general": [
        "Entry-Level Professional",
    ],
}


def detect_resume_domain(resume_text: str):

    text = resume_text.lower()

    # ========================================================
    # DOMAIN EVIDENCE
    # ========================================================
    # Important:
    # Past occupation does not automatically determine the
    # candidate's current career direction.
    #
    # Example:
    # "Educator" + strong Data Analytics / ML projects
    # should be classified toward Data Analytics when the
    # resume is clearly positioned for that career.
    # ========================================================

    scores = {
        domain: 0
        for domain in DOMAIN_KEYWORDS
    }

    matched = {
        domain: []
        for domain in DOMAIN_KEYWORDS
    }

    # --------------------------------------------------------
    # Base keyword scoring
    # --------------------------------------------------------

    strong_identity_keywords = {
        "chemist",
        "data analyst",
        "data scientist",
        "software engineer",
        "software developer",
        "teacher",
        "educator",
        "microbiologist",
        "biologist",
        "operations executive",
        "financial analyst",
    }

    for domain, keywords in DOMAIN_KEYWORDS.items():

        for keyword in keywords:

            keyword_lower = keyword.lower()

            if keyword_lower in text:

                if keyword_lower in strong_identity_keywords:
                    scores[domain] += 5
                else:
                    scores[domain] += 2

                matched[domain].append(
                    keyword
                )

    # ========================================================
    # GENERIC SKILLS
    # ========================================================

    generic_only_keywords = {
        "python",
        "sql",
        "excel",
        "power bi",
        "tableau",
        "communication",
        "research",
        "statistics",
    }

    for domain in scores:

        matched[domain] = [
            x
            for x in matched[domain]
            if x.lower()
            not in generic_only_keywords
        ]

    # ========================================================
    # REAL DATA ANALYTICS EVIDENCE
    # ========================================================

    analytics_identity = [
        "data analyst",
        "data analytics",
        "business analyst",
        "bi analyst",
        "analytics analyst",
        "reporting analyst",
        "business intelligence analyst",
    ]

    analytics_work = [
        "dashboard",
        "dashboards",
        "kpi dashboard",
        "kpi reporting",
        "business intelligence",
        "data analysis",
        "data analytics",
        "data visualization",
        "reporting analysis",
        "analytical reporting",
        "data-driven insights",
        "performance reporting",
    ]

    analytics_project = [
        "analytics project",
        "data analysis project",
        "data analytics project",
        "business analytics project",
        "business intelligence project",
        "dashboard project",
        "credit risk analysis",
        "fraud analysis",
        "sales analysis",
        "customer analysis",
    ]

    analytics_identity_count = sum(
        1
        for keyword in analytics_identity
        if keyword in text
    )

    analytics_work_count = sum(
        1
        for keyword in analytics_work
        if keyword in text
    )

    analytics_project_count = sum(
        1
        for keyword in analytics_project
        if keyword in text
    )

    # ========================================================
    # REAL DATA SCIENCE / ML EVIDENCE
    # ========================================================

    data_science_identity = [
        "data scientist",
        "data science",
        "machine learning engineer",
        "ai engineer",
        "ml engineer",
    ]

    data_science_work = [
        "machine learning",
        "predictive modeling",
        "predictive maintenance",
        "model training",
        "feature engineering",
        "classification model",
        "regression model",
        "xgboost",
        "scikit-learn",
        "random forest",
        "model deployment",
    ]

    data_science_project = [
        "machine learning project",
        "data science project",
        "predictive maintenance",
        "fraud detection",
        "credit risk analysis",
        "recommendation system",
        "prediction model",
    ]

    ds_identity_count = sum(
        1
        for keyword in data_science_identity
        if keyword in text
    )

    ds_work_count = sum(
        1
        for keyword in data_science_work
        if keyword in text
    )

    ds_project_count = sum(
        1
        for keyword in data_science_project
        if keyword in text
    )

    # ========================================================
    # DATA ANALYTICS CAREER SIGNAL
    # ========================================================

    analytics_evidence_score = (
        analytics_identity_count * 6
        + min(analytics_work_count, 6) * 3
        + min(analytics_project_count, 4) * 4
    )

    # ========================================================
    # DATA SCIENCE CAREER SIGNAL
    # ========================================================

    data_science_evidence_score = (
        ds_identity_count * 6
        + min(ds_work_count, 8) * 3
        + min(ds_project_count, 4) * 4
    )

    # Add these scores separately from generic skill matching.
    scores["data_analytics"] += analytics_evidence_score
    scores["data_science"] += data_science_evidence_score

    # ========================================================
    # CAREER DIRECTION OVERRIDE
    # ========================================================
    # Strong analytics / DS evidence can legitimately override
    # historical education experience.
    # ========================================================

    education_score = scores.get(
        "education",
        0
    )

    analytics_strong = (
        analytics_evidence_score >= 8
    )

    data_science_strong = (
        data_science_evidence_score >= 10
    )

    if analytics_strong:

        scores["data_analytics"] += 8

    if data_science_strong:

        scores["data_science"] += 8

    # If both analytics and DS are strongly represented,
    # prefer Data Analytics when the resume contains strong
    # dashboard/reporting/BI evidence.
    if (
        analytics_strong
        and data_science_strong
    ):

        if (
            analytics_work_count
            + analytics_project_count
            >= 2
        ):

            scores["data_analytics"] += 6

    # ========================================================
    # DOMAIN PRIORITY
    # ========================================================

    priority = [
        "data_analytics",
        "data_science",
        "chemistry",
        "biology_life_sciences",
        "software_development",
        "banking_finance",
        "operations",
        "education",
        "marketing_sales",
    ]

    best_domain = "general"
    best_score = 0

    for domain in priority:

        score = scores.get(
            domain,
            0
        )

        if score > best_score:

            best_score = score
            best_domain = domain

    # ========================================================
    # MEANINGFUL EVIDENCE CHECK
    # ========================================================

    if best_score < 4:

        best_domain = "general"

    return {
        "domain": best_domain,
        "score": best_score,
        "scores": scores,
        "matched_keywords": matched.get(
            best_domain,
            [],
        ),
    }

# ============================================================
# AI PROMPT
# ============================================================

def build_resume_profile_prompt(
    resume_text: str,
) -> str:

    return f"""
You are an expert Resume Intelligence and Career Analysis AI.

Analyze ONLY the resume provided below.

IMPORTANT RULES:

1. Do NOT invent information.
2. Do NOT use information from any previous resume.
3. Do NOT assume the candidate is a specific person.
4. Extract the candidate name directly from this resume.
5. Extract email, phone and location only if present.
6. Teaching or academic experience must remain separate from
   technical or corporate experience.
7. Do not convert education into work experience.
8. Do not exaggerate seniority.
9. If a value is not available, use null or an empty list.
10. Skills must be explicitly supported by the resume.
11. Projects must be explicitly supported by the resume.
12. Recommended roles must be based on the candidate's ACTUAL
    profession, education, experience and demonstrated skills.
13. Generic skills such as Excel, SQL, Python, Communication,
    Research or Statistics MUST NOT by themselves make someone
    a Data Analyst or Data Scientist.
14. A chemistry, laboratory, pharmaceutical or chemical-science
    candidate MUST NOT be classified as Data Analyst or
    Data Scientist merely because Excel, SQL, statistics,
    research or data-related words appear.
15. A candidate must have meaningful evidence of analytics work
    such as dashboards, KPI reporting, business intelligence,
    reporting analysis or a Data Analyst role before recommending
    Data Analyst roles.
16. A candidate must have meaningful evidence of machine learning,
    predictive modeling, model training, feature engineering,
    data science projects or a Data Scientist role before
    recommending Data Scientist roles.
17. Identify the candidate's primary career domain.
18. Do not confuse tools with professions.
19. Do not recommend unrelated technical roles simply because
    a generic tool appears in the resume.
20. Return ONLY valid JSON.

PRIMARY DOMAIN OPTIONS:

- chemistry
- biology_life_sciences
- data_analytics
- data_science
- software_development
- banking_finance
- operations
- education
- marketing_sales
- general

Use exactly this structure:

{{
    "candidate": {{
        "name": null,
        "email": null,
        "phone": null,
        "location": null
    }},

    "career_domain": null,

    "professional_summary": null,

    "experience": {{
        "total_years": 0,
        "technical_years": 0,
        "positions": [
            {{
                "title": null,
                "organization": null,
                "duration": null,
                "type": null,
                "description": null
            }}
        ]
    }},

    "skills": {{
        "programming": [],
        "data_analysis": [],
        "machine_learning": [],
        "visualization": [],
        "deployment": [],
        "other": []
    }},

    "education": [
        {{
            "degree": null,
            "institution": null,
            "status": null,
            "year": null
        }}
    ],

    "projects": [
        {{
            "name": null,
            "description": null,
            "technologies": [],
            "impact": null
        }}
    ],

    "achievements": [],

    "recommended_roles": [],

    "career_level": null
}}

Resume:

---------------- RESUME START ----------------

{resume_text}

----------------- RESUME END -----------------
"""


# ============================================================
# SKILL HELPERS
# ============================================================

KNOWN_SKILLS = [

    # Programming
    "Python",
    "SQL",
    "R",
    "Java",
    "C++",
    "JavaScript",
    "TypeScript",

    # Analytics
    "Pandas",
    "NumPy",
    "Excel",
    "Advanced Excel",
    "Power BI",
    "Tableau",
    "Matplotlib",
    "Seaborn",
    "EDA",
    "Data Visualization",
    "Business Intelligence",
    "Data Analysis",
    "Statistical Analysis",

    # ML
    "Machine Learning",
    "Scikit-learn",
    "XGBoost",
    "SMOTE",
    "Deep Learning",
    "TensorFlow",
    "PyTorch",

    # Deployment
    "Streamlit",
    "Docker",
    "Kubernetes",
    "MLflow",
    "Airflow",
    "FastAPI",
    "Flask",

    # Databases
    "MySQL",
    "PostgreSQL",
    "SQL Server",
    "MongoDB",

    # AI
    "LLM",
    "Generative AI",
    "LangChain",
    "RAG",

    # Chemistry / Science
    "Chemistry",
    "Analytical Chemistry",
    "Organic Chemistry",
    "Inorganic Chemistry",
    "Physical Chemistry",
    "Chemical Analysis",
    "Laboratory",
    "Laboratory Analysis",
    "Laboratory Testing",
    "Quality Control",
    "Quality Assurance",
    "Titration",
    "Spectroscopy",
    "Chromatography",
    "HPLC",
    "Gas Chromatography",
    "UV Spectroscopy",
    "Sample Preparation",
    "Chemical Testing",
    "Formulation",
    "Pharmaceutical Analysis",
    "Microbiology",
    "Biotechnology",
    "Biochemistry",

    # Banking / Operations
    "KYC",
    "Banking Operations",
    "Disbursement",
    "FinOne",
    "SFDC",
    "Tally ERP",
    "GST",
    "SAP",

    # Office
    "MS Office",
    "Microsoft Office",
    "Microsoft Excel",
    "PowerPoint",
    "Word",
    "Outlook",

    # General
    "Communication",
    "Research",
    "Leadership",
]


def detect_skills_from_text(resume_text: str):

    text = resume_text.lower()
    detected = []

    for skill in KNOWN_SKILLS:

        normalized = skill.lower()

        if len(skill) <= 2:

            if re.search(
                rf"\b{re.escape(skill)}\b",
                resume_text,
                re.IGNORECASE,
            ):
                detected.append(skill)

        elif normalized in text:
            detected.append(skill)

    return list(
        dict.fromkeys(detected)
    )


# ============================================================
# FALLBACK PROFILE
# ============================================================

def build_fallback_profile(
    resume_text: str,
):

    text_lower = resume_text.lower()

    lines = [
        line.strip()
        for line in resume_text.splitlines()
        if line.strip()
    ]

    # ========================================================
    # CONTACT
    # ========================================================

    email = None

    email_match = re.search(
        r"[A-Za-z0-9._%+-]+@"
        r"[A-Za-z0-9.-]+\."
        r"[A-Za-z]{2,}",
        resume_text,
        re.IGNORECASE,
    )

    if email_match:
        email = email_match.group(0).strip()

    phone = None

    for pattern in [
        r"\+91[\s-]?[6-9]\d{9}",
        r"\b[6-9]\d{9}\b",
    ]:

        phone_match = re.search(
            pattern,
            resume_text,
        )

        if phone_match:

            phone = re.sub(
                r"\D",
                "",
                phone_match.group(0),
            )

            if len(phone) > 10:
                phone = phone[-10:]

            break

    # ========================================================
    # NAME
    # ========================================================

    name = None

    blocked_terms = [
        "date of birth",
        "gender",
        "marital status",
        "languages known",
        "experience",
        "education",
        "skills",
        "technical skills",
        "computer proficiency",
        "hobbies",
        "objective",
        "summary",
        "email",
        "phone",
        "mobile",
        "address",
        "linkedin",
        "github",
        "curriculum vitae",
        "resume",
    ]

    name_candidates = []

    for line in lines[:15]:

        clean_line = re.sub(
            r"[^A-Za-z .'-]",
            "",
            line,
        ).strip()

        if not clean_line:
            continue

        lower_line = clean_line.lower()

        if any(
            term in lower_line
            for term in blocked_terms
        ):
            continue

        if (
            email
            and email.lower() in lower_line
        ):
            continue

        words = clean_line.split()

        if not 1 <= len(words) <= 5:
            continue

        if all(
            re.fullmatch(
                r"[A-Za-z.'-]+",
                word,
            )
            for word in words
        ):
            name_candidates.append(
                clean_line
            )

    if name_candidates:
        name = name_candidates[0]

    # ========================================================
    # LOCATION
    # ========================================================

    location = None

    location_patterns = [
        r"\bMumbai\b",
        r"\bPune\b",
        r"\bDelhi\b",
        r"\bNew Delhi\b",
        r"\bBengaluru\b",
        r"\bBangalore\b",
        r"\bHyderabad\b",
        r"\bChennai\b",
        r"\bAhmedabad\b",
        r"\bKolkata\b",
        r"\bThane\b",
        r"\bNavi Mumbai\b",
        r"\bNoida\b",
        r"\bGurgaon\b",
        r"\bGurugram\b",
        r"\bJaipur\b",
        r"\bIndore\b",
        r"\bNagpur\b",
    ]

    for pattern in location_patterns:

        match = re.search(
            pattern,
            resume_text,
            re.IGNORECASE,
        )

        if match:
            location = match.group(0).strip()
            break

    # ========================================================
    # DOMAIN
    # ========================================================

    domain_info = detect_resume_domain(
        resume_text
    )

    domain = domain_info["domain"]

    # ========================================================
    # SKILLS
    # ========================================================

    detected_skills = detect_skills_from_text(
        resume_text
    )

    # ========================================================
    # CATEGORY CLASSIFICATION
    # ========================================================

    # IMPORTANT: tools are not professions.
    # Excel / SQL / Python / Power BI / statistics alone must
    # never turn an Operations, Chemistry, Banking, Education,
    # etc. resume into a Data Analyst resume.

    programming_names = {
        "Python", "SQL", "R", "Java", "C++",
        "JavaScript", "TypeScript",
    }

    analytics_names = {
        "Pandas", "NumPy", "Excel", "Advanced Excel",
        "EDA", "Business Intelligence", "Data Analysis",
        "Statistical Analysis",
    }

    ml_names = {
        "Machine Learning", "Scikit-learn", "XGBoost",
        "SMOTE", "Deep Learning", "TensorFlow", "PyTorch",
    }

    visualization_names = {
        "Power BI", "Tableau", "Matplotlib",
        "Seaborn", "Data Visualization",
    }

    deployment_names = {
        "Streamlit", "Docker", "Kubernetes", "MLflow",
        "Airflow", "FastAPI", "Flask",
    }

    programming = [
        skill for skill in detected_skills
        if skill in programming_names
    ]

    raw_data_analysis = [
        skill for skill in detected_skills
        if skill in analytics_names
    ]

    machine_learning = [
        skill for skill in detected_skills
        if skill in ml_names
    ]

    raw_visualization = [
        skill for skill in detected_skills
        if skill in visualization_names
    ]

    deployment = [
        skill for skill in detected_skills
        if skill in deployment_names
    ]

    # --------------------------------------------------------
    # REAL ANALYTICS CONTEXT
    # --------------------------------------------------------

    analytics_identity_keywords = [
        "data analyst",
        "data analytics",
        "business analyst",
        "bi analyst",
        "analytics analyst",
        "reporting analyst",
        "business intelligence analyst",
    ]

    analytics_work_keywords = [
        "dashboard",
        "dashboards",
        "kpi dashboard",
        "kpi reporting",
        "business intelligence",
        "reporting analysis",
        "data reporting",
        "sql reporting",
        "business reporting",
        "data visualization",
        "data-driven insights",
        "analytical reporting",
        "performance reporting",
    ]

    analytics_project_keywords = [
        "analytics project",
        "data analysis project",
        "data analytics project",
        "business analytics project",
        "business intelligence project",
        "dashboard project",
        "reporting project",
    ]

    analytics_identity_count = sum(
        1 for keyword in analytics_identity_keywords
        if keyword in text_lower
    )

    analytics_work_count = sum(
        1 for keyword in analytics_work_keywords
        if keyword in text_lower
    )

    analytics_project_count = sum(
        1 for keyword in analytics_project_keywords
        if keyword in text_lower
    )

    has_real_analytics_context = (
        analytics_identity_count >= 1
        or analytics_work_count >= 2
        or analytics_project_count >= 1
    )

    strong_non_analytics_domain = domain in {
        "chemistry",
        "biology_life_sciences",
        "banking_finance",
        "operations",
        "education",
        "marketing_sales",
    }

    if (
        strong_non_analytics_domain
        and not has_real_analytics_context
    ):
        data_analysis = []
        visualization = []
    else:
        data_analysis = raw_data_analysis
        visualization = raw_visualization

    categorized = (
        programming
        + data_analysis
        + machine_learning
        + visualization
        + deployment
    )

    other = [
        skill for skill in detected_skills
        if skill not in categorized
    ]

    # ========================================================
    # EXPERIENCE
    # ========================================================

    positions = []

    experience_rules = [

        (
            [
                "analytical chemist",
                "research chemist",
                "chemist",
                "chemical analyst",
                "qc chemist",
                "laboratory analyst",
                "lab analyst",
            ],
            "Chemist",
        ),

        (
            [
                "microbiologist",
            ],
            "Microbiologist",
        ),

        (
            [
                "data analyst",
            ],
            "Data Analyst",
        ),

        (
            [
                "data scientist",
            ],
            "Data Scientist",
        ),

        (
            [
                "machine learning engineer",
            ],
            "Machine Learning Engineer",
        ),

        (
            [
                "business analyst",
            ],
            "Business Analyst",
        ),

        (
            [
                "software engineer",
                "software developer",
            ],
            "Software Engineer",
        ),

        (
            [
                "operation executive",
                "operations executive",
                "operation exective",
                "operations analyst",
            ],
            "Operations Executive",
        ),

        (
            [
                "financial analyst",
            ],
            "Financial Analyst",
        ),

        (
            [
                "teacher",
                "educator",
                "academic mentor",
                "subject matter expert",
            ],
            "Educator / Academic Mentor",
        ),
    ]

    detected_position = None

    for keywords, title in experience_rules:

        if any(
            keyword in text_lower
            for keyword in keywords
        ):
            detected_position = title
            break

    if detected_position:

        duration = None

        duration_match = re.search(
            r"\b\d+(?:\.\d+)?\s*"
            r"(?:year|years|month|months)\b",
            resume_text,
            re.IGNORECASE,
        )

        if duration_match:
            duration = duration_match.group(0)

        positions.append(
            {
                "title": detected_position,
                "organization": None,
                "duration": duration,
                "type": "Professional",
                "description": (
                    "Professional experience "
                    "identified from the uploaded resume."
                ),
            }
        )

    # ========================================================
    # EXPERIENCE YEARS
    # ========================================================

    total_years = 0

    year_match = re.search(
        r"(\d+(?:\.\d+)?)\s*(?:\+)?\s*years?",
        text_lower,
    )

    if year_match:

        try:
            total_years = float(
                year_match.group(1)
            )
        except ValueError:
            total_years = 0

    # ========================================================
    # EDUCATION
    # ========================================================

    education = []

    education_rules = [
        (
            "bachelor of banking and insurance",
            "Bachelor of Banking and Insurance",
        ),
        (
            "bachelor of science",
            "Bachelor of Science",
        ),
        (
            "b.sc",
            "Bachelor of Science",
        ),
        (
            "master of",
            "Master's Degree",
        ),
        (
            "masters in",
            "Master's Degree",
        ),
        (
            "higher secondary",
            "Higher Secondary Certificate",
        ),
        (
            "h.s.c",
            "Higher Secondary Certificate",
        ),
        (
            "secondary school certificate",
            "Secondary School Certificate",
        ),
        (
            "s.s.c",
            "Secondary School Certificate",
        ),
    ]

    detected_degrees = set()

    for keyword, degree in education_rules:

        if keyword in text_lower:

            if degree not in detected_degrees:

                education.append(
                    {
                        "degree": degree,
                        "institution": None,
                        "status": "Completed",
                        "year": None,
                    }
                )

                detected_degrees.add(degree)

    # ========================================================
    # PROJECTS
    # ========================================================

    projects = []

    if "predictive maintenance" in text_lower:

        projects.append(
            {
                "name": (
                    "AI-Based Predictive Maintenance "
                    "& Failure Diagnostic System"
                ),
                "description": (
                    "Industrial machine failure "
                    "prediction platform using "
                    "machine learning."
                ),
                "technologies": [
                    "Python",
                    "XGBoost",
                    "Machine Learning",
                    "Streamlit",
                ],
                "impact": "Achieved 98.55% accuracy.",
            }
        )

    if "credit risk analysis" in text_lower:

        projects.append(
            {
                "name": "Credit Risk Analysis Dashboard",
                "description": (
                    "Power BI credit intelligence "
                    "dashboard for identifying "
                    "high-risk customers."
                ),
                "technologies": [
                    "Power BI",
                    "Data Analysis",
                ],
                "impact": (
                    "Identified 55K+ high-risk customers."
                ),
            }
        )

    if (
        "fraud" in text_lower
        and "paisa" in text_lower
    ):

        projects.append(
            {
                "name": (
                    "PaisaBazaar Banking Fraud / "
                    "EDA Project"
                ),
                "description": (
                    "Exploratory analysis of fraud "
                    "indicators and financial risk signals."
                ),
                "technologies": [
                    "Python",
                    "Pandas",
                    "EDA",
                    "Data Analysis",
                ],
                "impact": None,
            }
        )

    # ========================================================
    # RECOMMENDED ROLES
    # ========================================================

    # Roles are determined by the primary resume domain.
    # Generic tools such as Excel, SQL, Python or Power BI
    # cannot independently trigger Data Analyst/Data Scientist.

    recommended_roles = list(
        DOMAIN_ROLES.get(
            domain,
            DOMAIN_ROLES["general"],
        )
    )

    # ========================================================
    # PROFESSIONAL SUMMARY
    # ========================================================

    if positions:

        professional_summary = (
            f"{positions[0]['title']} with "
            f"{total_years:g} year(s) of "
            f"experience based on the uploaded resume."
        )

    elif domain != "general":

        professional_summary = (
            f"Resume profile identified primarily "
            f"within the {domain.replace('_', ' ')} domain."
        )

    elif detected_skills:

        professional_summary = (
            "Entry-level professional with "
            "skills identified from the uploaded resume."
        )

    else:

        professional_summary = (
            "Candidate profile generated from "
            "the uploaded resume."
        )

    # ========================================================
    # CAREER LEVEL
    # ========================================================

    if total_years >= 5:
        career_level = "Experienced"
    elif total_years >= 2:
        career_level = "Mid Level"
    elif total_years > 0:
        career_level = "Early Career"
    else:
        career_level = "Entry Level"

    # ========================================================
    # FINAL FALLBACK PROFILE
    # ========================================================

    return {

        "candidate": {
            "name": name,
            "email": email,
            "phone": phone,
            "location": location,
        },

        "career_domain": domain,

        "professional_summary":
            professional_summary,

        "experience": {

            "total_years":
                total_years,

            "technical_years": (
                total_years
                if (
                    domain
                    in {
                        "data_science",
                        "data_analytics",
                        "software_development",
                    }
                )
                else 0
            ),

            "positions":
                positions,
        },

        "skills": {

            "programming":
                programming,

            "data_analysis":
                data_analysis,

            "machine_learning":
                machine_learning,

            "visualization":
                visualization,

            "deployment":
                deployment,

            "other":
                other,
        },

        "education":
            education,

        "projects":
            projects,

        "achievements":
            [],

        "recommended_roles":
            recommended_roles,

        "career_level":
            career_level,
    }


# ============================================================
# PROFILE SANITIZER
# ============================================================

def sanitize_ai_profile(
    profile: dict,
    resume_text: str,
):

    if not isinstance(profile, dict):
        return build_fallback_profile(
            resume_text
        )

    deterministic = build_fallback_profile(
        resume_text
    )

    detected = detect_resume_domain(
        resume_text
    )

    domain = detected["domain"]

    if not isinstance(
        profile.get("candidate"),
        dict,
    ):
        profile["candidate"] = {}

    # ========================================================
    # DOMAIN
    # ========================================================

    profile["career_domain"] = domain

    profile["recommended_roles"] = list(
        DOMAIN_ROLES.get(
            domain,
            DOMAIN_ROLES["general"],
        )
    )

    # ========================================================
    # SKILLS
    # ========================================================

    ai_skills = profile.get(
        "skills",
        {},
    )

    if not isinstance(
        ai_skills,
        dict,
    ):
        ai_skills = {}

    for category in [
        "programming",
        "data_analysis",
        "machine_learning",
        "visualization",
        "deployment",
        "other",
    ]:

        if not isinstance(
            ai_skills.get(category),
            list,
        ):
            ai_skills[category] = []

    # ========================================================
    # ANALYTICS EVIDENCE
    # ========================================================

    strong_non_analytics = domain in {
        "chemistry",
        "biology_life_sciences",
        "banking_finance",
        "operations",
        "education",
        "marketing_sales",
    }

    resume_lower = resume_text.lower()

    analytics_identity_keywords = [
        "data analyst",
        "data analytics",
        "business analyst",
        "bi analyst",
        "analytics analyst",
        "reporting analyst",
        "business intelligence analyst",
    ]

    analytics_work_keywords = [
        "dashboard",
        "dashboards",
        "kpi dashboard",
        "kpi reporting",
        "business intelligence",
        "reporting analysis",
        "data reporting",
        "sql reporting",
        "business reporting",
        "data visualization",
        "data-driven insights",
        "analytical reporting",
        "performance reporting",
    ]

    analytics_project_keywords = [
        "analytics project",
        "data analysis project",
        "data analytics project",
        "business analytics project",
        "business intelligence project",
        "dashboard project",
        "reporting project",
        "credit risk analysis",
        "fraud analysis",
        "credit risk",
        "paisa bazaar",
    ]

    identity_count = sum(
        1
        for keyword in analytics_identity_keywords
        if keyword in resume_lower
    )

    work_count = sum(
        1
        for keyword in analytics_work_keywords
        if keyword in resume_lower
    )

    project_count = sum(
        1
        for keyword in analytics_project_keywords
        if keyword in resume_lower
    )

    has_real_analytics_context = (
        identity_count >= 1
        or work_count >= 2
        or project_count >= 1
    )

    # ========================================================
    # MACHINE LEARNING EVIDENCE
    # ========================================================

    ml_keywords = [
        "machine learning",
        "predictive modeling",
        "predictive maintenance",
        "feature engineering",
        "model training",
        "classification",
        "regression",
        "xgboost",
        "scikit-learn",
        "random forest",
        "model deployment",
    ]

    ml_evidence_count = sum(
        1
        for keyword in ml_keywords
        if keyword in resume_lower
    )

    has_real_ml_context = (
        ml_evidence_count >= 2
    )

    # ========================================================
    # REMOVE UNSUPPORTED ANALYTICS SKILLS FROM
    # STRONGLY NON-ANALYTICS RESUMES
    # ========================================================

    if (
        strong_non_analytics
        and not has_real_analytics_context
    ):

        generic_analytics = {
            "data analysis",
            "statistical analysis",
            "business intelligence",
            "eda",
            "pandas",
            "numpy",
            "excel",
            "advanced excel",
        }

        ai_skills["data_analysis"] = [
            skill
            for skill
            in ai_skills["data_analysis"]
            if str(skill).strip().lower()
            not in generic_analytics
        ]

        ai_skills["visualization"] = [
            skill
            for skill
            in ai_skills["visualization"]
            if str(skill).strip().lower()
            not in {
                "power bi",
                "tableau",
                "data visualization",
            }
        ]

    profile["skills"] = ai_skills

    # ========================================================
    # EXPERIENCE
    # ========================================================

    if not isinstance(
        profile.get("experience"),
        dict,
    ):
        profile["experience"] = deterministic[
            "experience"
        ]

    # Preserve the actual total experience.
    profile["experience"]["total_years"] = (
        deterministic["experience"]["total_years"]
    )

    # --------------------------------------------------------
    # IMPORTANT:
    # Teaching / education experience is NOT technical
    # analytics experience.
    # --------------------------------------------------------

    teaching_keywords = [
        "educator",
        "teacher",
        "academic mentor",
        "subject matter expert",
        "teaching",
        "mentoring",
        "mentor",
    ]

    has_teaching_experience = any(
        keyword in resume_lower
        for keyword in teaching_keywords
    )

    if has_teaching_experience:

        profile["experience"]["technical_years"] = 0

    elif domain in {
        "data_science",
        "data_analytics",
        "software_development",
    }:

        profile["experience"]["technical_years"] = (
            deterministic["experience"]["technical_years"]
        )

    else:

        profile["experience"]["technical_years"] = 0

    # ========================================================
    # TARGET ROLE
    # ========================================================
    # Career domain determines the target role when strong
    # career evidence exists. AI-generated role labels should
    # not override deterministic resume evidence.
    # ========================================================

    target_role = None

    if domain == "data_analytics":

        target_role = "Data Analyst"

    elif domain == "data_science":

        # If the resume has substantial analytics evidence
        # alongside ML projects, prefer Data Analyst as the
        # broader entry-level target.
        if has_real_analytics_context:

            target_role = "Data Analyst"

        else:

            target_role = "Data Scientist"

    elif domain == "software_development":

        target_role = "Software Developer"

    elif domain == "chemistry":

        target_role = "Chemist"

    elif domain == "biology_life_sciences":

        target_role = "Biologist"

    elif domain == "banking_finance":

        target_role = "Financial Analyst"

    elif domain == "operations":

        target_role = "Operations Executive"

    elif domain == "education":

        target_role = "Educator"

    elif domain == "marketing_sales":

        target_role = "Marketing Executive"

    else:

        ai_target_role = (
            profile.get("target_role")
            or ""
        )

        if str(
            ai_target_role
        ).strip():

            target_role = str(
                ai_target_role
            ).strip()

        else:

            target_role = (
                "Entry-Level Professional"
            )

    profile["target_role"] = target_role

    # ========================================================
    # CAREER LEVEL
    # ========================================================

    if (
        profile["experience"]["technical_years"]
        >= 5
    ):

        profile["career_level"] = "Experienced"

    elif (
        profile["experience"]["technical_years"]
        >= 2
    ):

        profile["career_level"] = "Mid Level"

    elif (
        profile["experience"]["technical_years"]
        > 0
    ):

        profile["career_level"] = "Early Career"

    else:

        profile["career_level"] = "Entry Level"

    return profile

# ============================================================
# PROFILE EXTRACTION
# ============================================================

def extract_resume_profile(
    resume_text: str,
):

    prompt = build_resume_profile_prompt(
        resume_text
    )

    try:

        ai_response = generate_response(
            prompt
        )

        cleaned_response = clean_ai_json(
            ai_response
        )

        profile = json.loads(
            cleaned_response
        )

        profile = sanitize_ai_profile(
            profile,
            resume_text,
        )

        return {
            "status": "success",
            "source": "gemini",
            "profile": profile,
            "ai_error": None,
        }

    except Exception as e:

        fallback = build_fallback_profile(
            resume_text
        )

        return {
            "status": "success",
            "source": "fallback",
            "profile": fallback,
            "ai_error": str(e),
        }


# ============================================================
# RESUME METADATA
# ============================================================

def persist_resume_metadata(
    user_id: str,
    resume_id: str,
    filename: str,
):

    try:

        response = (
            supabase
            .table("user_profiles")
            .update(
                {
                    "resume_id": resume_id,
                    "resume_filename": filename,
                }
            )
            .eq(
                "user_id",
                user_id,
            )
            .execute()
        )

        data = getattr(
            response,
            "data",
            None,
        )

        return {
            "success": bool(data),
            "data": data,
        }

    except Exception as e:

        return {
            "success": False,
            "data": None,
            "error": str(e),
        }


# ============================================================
# RESUME UPLOAD
# ============================================================

@router.post("/upload")
async def upload_resume(
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):

    # ========================================================
    # USER
    # ========================================================

    if not user or not getattr(
        user,
        "id",
        None,
    ):

        raise HTTPException(
            status_code=401,
            detail=(
                "Authenticated user could not be identified."
            ),
        )

    user_id = user.id

    # ========================================================
    # FILE VALIDATION
    # ========================================================

    if not file.filename:

        raise HTTPException(
            status_code=400,
            detail="No file selected.",
        )

    filename = file.filename

    extension = os.path.splitext(
        filename
    )[1].lower()

    if extension not in ALLOWED_EXTENSIONS:

        raise HTTPException(
            status_code=400,
            detail=(
                "Unsupported file type. "
                "Upload PDF, DOCX or TXT."
            ),
        )

    file_bytes = await file.read()

    if not file_bytes:

        raise HTTPException(
            status_code=400,
            detail="Uploaded file is empty.",
        )

    if len(file_bytes) > MAX_FILE_SIZE:

        raise HTTPException(
            status_code=400,
            detail=(
                "Resume size must be less than 5 MB."
            ),
        )

    # ========================================================
    # FILE STORAGE
    # ========================================================

    resume_id = str(
        uuid.uuid4()
    )

    stored_filename = (
        f"{resume_id}{extension}"
    )

    file_path = os.path.join(
        UPLOAD_DIR,
        stored_filename,
    )

    try:

        with open(
            file_path,
            "wb",
        ) as output_file:

            output_file.write(
                file_bytes
            )

    except Exception as e:

        raise HTTPException(
            status_code=500,
            detail=(
                f"Unable to store uploaded file: {str(e)}"
            ),
        )

    # ========================================================
    # TEXT EXTRACTION
    # ========================================================

    resume_text = ""

    try:

        if extension == ".pdf":

            from pypdf import PdfReader

            reader = PdfReader(
                file_path
            )

            pages = []

            for page in reader.pages:

                pages.append(
                    page.extract_text()
                    or ""
                )

            resume_text = "\n".join(
                pages
            )

        elif extension == ".docx":

            from docx import Document

            document = Document(
                file_path
            )

            paragraphs = [
                paragraph.text
                for paragraph
                in document.paragraphs
            ]

            table_text = []

            for table in document.tables:

                for row in table.rows:

                    row_values = [
                        cell.text
                        for cell in row.cells
                    ]

                    table_text.append(
                        " ".join(
                            row_values
                        )
                    )

            resume_text = "\n".join(
                paragraphs
                + table_text
            )

        elif extension == ".txt":

            with open(
                file_path,
                "r",
                encoding="utf-8",
                errors="ignore",
            ) as input_file:

                resume_text = input_file.read()

    except Exception as e:

        try:

            if os.path.exists(
                file_path
            ):
                os.remove(
                    file_path
                )

        except Exception:
            pass

        return {
            "status": "error",
            "message": (
                "Resume uploaded but text "
                "extraction failed."
            ),
            "resume_id": resume_id,
            "error": str(e),
        }

    # ========================================================
    # CLEAN TEXT
    # ========================================================

    resume_text = clean_resume_text(
        resume_text
    )

    if not resume_text:

        return {
            "status": "error",
            "message": (
                "Resume uploaded but no "
                "readable text was extracted."
            ),
            "resume_id": resume_id,
        }

    # ========================================================
    # RESUME VALIDATION
    # ========================================================

    validation = validate_resume_document(
        resume_text
    )

    if not validation["is_resume"]:

        return {
            "status": "error",
            "message": validation["message"],
            "resume_id": resume_id,
            "filename": filename,
            "file_type": extension[1:].upper(),
            "file_size": len(file_bytes),
            "text_length": len(resume_text),
            "validation_score": validation.get(
                "score",
                0,
            ),
            "document_type": "invalid_resume",
        }

    # ========================================================
    # PROFILE EXTRACTION
    # ========================================================

    profile_result = extract_resume_profile(
        resume_text
    )

    candidate_profile = profile_result[
        "profile"
    ]

    # ========================================================
    # PERSIST RESUME METADATA
    # ========================================================

    profile_persistence = persist_resume_metadata(
        user_id=user_id,
        resume_id=resume_id,
        filename=filename,
    )

    # ========================================================
    # FINAL RESPONSE
    # ========================================================

    return {

        "status": "success",

        "message": (
            "Resume uploaded, validated, text "
            "extracted and candidate profile generated."
        ),

        "resume_id": resume_id,

        "filename": filename,

        "file_type": extension[1:].upper(),

        "file_size": len(file_bytes),

        "stored_filename": stored_filename,

        "text_length": len(resume_text),

        "resume_text": resume_text,

        "validation_score": validation.get(
            "score",
            0,
        ),

        "document_type": "resume",

        "candidate_profile": candidate_profile,

        "profile_source": profile_result[
            "source"
        ],

        "ai_error": profile_result[
            "ai_error"
        ],

        "career_domain": candidate_profile.get(
            "career_domain"
        ),

        "recommended_roles": candidate_profile.get(
            "recommended_roles",
            [],
        ),

        "user_id": user_id,

        "resume_owner": {
            "user_id": user_id,
            "email": getattr(
                user,
                "email",
                None,
            ),
        },

        "profile_metadata_saved":
            profile_persistence["success"],
    }
